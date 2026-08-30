import os
import re
import requests
import numpy as np
import streamlit as st
import pymupdf
import pytesseract

from PIL import Image, ImageDraw
from docx import Document



# ============================================================
# CONFIGURATION
# ============================================================

OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "qwen2.5:1.5b"

EMBEDDING_MODEL = "all-MiniLM-L6-v2"

TEMP_DIR = "temp"
PAGES_DIR = os.path.join(TEMP_DIR, "pages")


# ============================================================
# STREAMLIT PAGE CONFIGURATION
# ============================================================

st.set_page_config(
    page_title="Offline Multimodal RAG",
    page_icon="📄",
    layout="wide"
)


# ============================================================
# POLISHED HEADER
# ============================================================

st.title("📄 Offline Multimodal RAG")

st.markdown(
    """
    <div style="
        padding: 14px 18px;
        border-radius: 12px;
        background: linear-gradient(
            90deg,
            #eef6ff,
            #f7fbff
        );
        border: 1px solid #d8e8f7;
        margin-bottom: 14px;
    ">
        <b>🔒 Private document intelligence</b><br>
        Ask questions about your documents and get
        answers grounded in retrieved evidence.
    </div>
    """,
    unsafe_allow_html=True
)

st.caption(
    "📄 PDF  •  📝 DOCX  •  🖼️ PNG/JPG  •  🔤 OCR"
)

st.markdown(
    """
    <div style="
        text-align: center;
        padding: 10px;
        margin-bottom: 20px;
        color: #555;
        font-size: 15px;
    ">
        📂 Upload
        &nbsp;→&nbsp;
        🔤 Extract / OCR
        &nbsp;→&nbsp;
        🔎 Retrieve
        &nbsp;→&nbsp;
        🤖 Local LLM
        &nbsp;→&nbsp;
        📌 Grounded Answer
    </div>
    """,
    unsafe_allow_html=True
)


# ============================================================
# LOAD EMBEDDING MODEL
# ============================================================

@st.cache_resource
def load_embedding_model():

    from sentence_transformers import SentenceTransformer

    return SentenceTransformer(
        EMBEDDING_MODEL
    )




# ============================================================
# TEXT UTILITIES
# ============================================================

def get_words(text):

    words = re.findall(
        r"\b[a-zA-Z0-9]+\b",
        text.lower()
    )

    return set(
        word
        for word in words
        if len(word) > 2
    )


def normalize_text(text):

    text = text.lower()

    text = re.sub(
        r"[^a-zA-Z0-9\s]",
        " ",
        text
    )

    text = re.sub(
        r"\s+",
        " ",
        text
    )

    return text.strip()


def split_into_units(text):

    units = []

    # --------------------------------------------------------
    # OCR / text lines
    # --------------------------------------------------------

    for line in text.splitlines():

        line = line.strip()

        if len(line) >= 3:

            units.append(line)


    # --------------------------------------------------------
    # Sentences
    # --------------------------------------------------------

    sentences = re.split(
        r"(?<=[.!?])\s+",
        text
    )

    for sentence in sentences:

        sentence = sentence.strip()

        if len(sentence) >= 10:

            units.append(sentence)


    # --------------------------------------------------------
    # Remove duplicates
    # --------------------------------------------------------

    result = []

    seen = set()

    for unit in units:

        key = unit.lower()

        if key not in seen:

            seen.add(key)

            result.append(unit)


    return result


# ============================================================
# LOCAL QWEN ANSWER GENERATION
# ============================================================

def generate_answer(question, evidence):

    prompt = f"""
You are a strict document question-answering assistant.

RULES:

1. Use ONLY the evidence provided below.
2. Do NOT use outside knowledge.
3. Do NOT guess.
4. Do NOT invent information.
5. If the answer is present, give the shortest direct answer.
6. If the answer is not present, say exactly:
Not found in the provided evidence.
7. Preserve numbers, IP addresses, commands, names,
   technical terms and values from the evidence.
8. Do not explain your reasoning.
9. Do not mention page numbers.
10. Do not add information that is not explicitly supported
    by the evidence.

Question:
{question}

Evidence:
{evidence}

Direct answer:
"""

    response = requests.post(

        OLLAMA_URL,

        json={
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0
            }
        },

        timeout=120
    )


    response.raise_for_status()


    data = response.json()


    return data.get(
        "response",
        ""
    ).strip()


# ============================================================
# PDF PROCESSOR
# ============================================================

def process_pdf(file_path):

    os.makedirs(
        PAGES_DIR,
        exist_ok=True
    )


    doc = pymupdf.open(
        file_path
    )


    pages = []


    for page_number, page in enumerate(doc):

        # ----------------------------------------------------
        # Native PDF text extraction
        # ----------------------------------------------------

        text = page.get_text(
            "text"
        ).strip()


        # ----------------------------------------------------
        # Render visual page
        # ----------------------------------------------------

        image_path = os.path.join(

            PAGES_DIR,

            f"page_{page_number + 1}.png"

        )


        pix = page.get_pixmap(

            matrix=pymupdf.Matrix(
                1.5,
                1.5
            )

        )


        pix.save(
            image_path
        )


        # ----------------------------------------------------
        # OCR fallback
        # ----------------------------------------------------

        if len(text) < 20:

            try:

                page_image = Image.open(
                    image_path
                )


                ocr_text = pytesseract.image_to_string(

                    page_image

                ).strip()


                if ocr_text:

                    text = ocr_text


            except Exception:

                pass


        pages.append({

            "page": page_number + 1,

            "text": text,

            "image": image_path

        })


    doc.close()


    return pages


# ============================================================
# DOCX PROCESSOR
# ============================================================

def process_docx(file_path):

    doc = Document(
        file_path
    )


    content = []


    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in doc.paragraphs:

        text = paragraph.text.strip()

        if text:

            content.append(
                text
            )


    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table in doc.tables:

        for row in table.rows:

            cells = []


            for cell in row.cells:

                text = cell.text.strip()

                if text:

                    cells.append(
                        text
                    )


            if cells:

                content.append(

                    " | ".join(
                        cells
                    )

                )


    if not content:

        return []


    text = "\n".join(
        content
    )


    os.makedirs(
        PAGES_DIR,
        exist_ok=True
    )


    # --------------------------------------------------------
    # DOCX visual preview
    # --------------------------------------------------------

    image_path = os.path.join(

        PAGES_DIR,

        "docx_preview.png"

    )


    width = 1200
    height = 800


    image = Image.new(

        "RGB",

        (
            width,
            height
        ),

        "white"

    )


    draw = ImageDraw.Draw(
        image
    )


    draw.text(

        (40, 40),

        "DOCX Document Preview",

        fill="black"

    )


    y = 100


    for line in content[:20]:

        draw.text(

            (40, y),

            line[:110],

            fill="black"

        )


        y += 30


        if y > 750:

            break


    image.save(
        image_path
    )


    return [

        {
            "page": 1,

            "text": text,

            "image": image_path

        }

    ]


# ============================================================
# IMAGE PROCESSOR
# ============================================================

def process_image(file_path):

    image = Image.open(
        file_path
    )


    image = image.convert(
        "RGB"
    )


    os.makedirs(
        PAGES_DIR,
        exist_ok=True
    )


    image_path = os.path.join(

        PAGES_DIR,

        "uploaded_image.png"

    )


    image.save(
        image_path
    )


    # --------------------------------------------------------
    # Tesseract OCR
    # --------------------------------------------------------

    ocr_text = pytesseract.image_to_string(

        image

    ).strip()


    return [

        {
            "page": 1,

            "text": ocr_text,

            "image": image_path

        }

    ]

# ============================================================
# FAST LOCAL RETRIEVAL
# ============================================================

def retrieve_evidence(
    question,
    text_pages
):

    """
    Fast local retrieval.

    Uses:
    - keyword overlap
    - exact phrase matching
    - factual-term matching
    - numeric/IP matching
    - page-level scoring

    No embedding model is loaded here.
    This keeps question answering fast.
    """

    # --------------------------------------------------------
    # Basic text collection
    # --------------------------------------------------------

    texts = [

        page.get("text", "")

        for page in text_pages

    ]


    # ========================================================
    # STOP WORDS
    # ========================================================

    stop_words = {

        "what",
        "which",
        "where",
        "when",
        "who",
        "why",
        "how",

        "is",
        "are",
        "was",
        "were",

        "does",
        "did",
        "do",

        "the",
        "a",
        "an",

        "this",
        "that",
        "these",
        "those",

        "show",
        "shown",

        "give",
        "given",

        "tell",

        "list",
        "find",

        "from",
        "document",
        "documents",
        "page",
        "pages",

        "please",

        "can",
        "could",
        "would",
        "should",

        "me",

        "of",
        "to",
        "in",
        "on",
        "for",
        "and",
        "or",

        "be",
        "been",
        "being",

        "about",

        "does",
        "did"

    }


    # ========================================================
    # QUESTION WORDS
    # ========================================================

    question_words = get_words(
        question
    )


    content_words = (

        question_words

        -

        stop_words

    )


    # ========================================================
    # FACTUAL TERMS
    # ========================================================

    factual_terms = {

        "ipv4",
        "ipv6",

        "address",
        "addresses",

        "gateway",
        "subnet",
        "mask",

        "dns",
        "server",
        "servers",

        "command",
        "commands",

        "port",
        "ports",

        "protocol",
        "protocols",

        "hostname",

        "adapter",
        "adapters",

        "interface",
        "interfaces",

        "username",
        "name",

        "date",
        "time",

        "version",

        "ip",

        "tcp",
        "udp",

        "tcpdump",
        "netstat",
        "ifconfig",
        "nslookup",
        "traceroute",
        "ping",

        "http",
        "https",

        "mac",

        "dhcp",

        "routing",

        "network",
        "networking"

    }


    requested_factual_terms = (

        content_words.intersection(
            factual_terms
        )

    )


    # ========================================================
    # NORMALIZED QUESTION
    # ========================================================

    question_normalized = normalize_text(
        question
    )


    # ========================================================
    # NUMBERS / IP ADDRESSES / TECHNICAL VALUES
    # ========================================================

    question_values = set(

        re.findall(

            r"\b\d+(?:\.\d+){0,3}\b",

            question

        )

    )


    # ========================================================
    # SCORE EVERY PAGE
    # ========================================================

    page_scores = []


    for page_text in texts:

        page_normalized = normalize_text(
            page_text
        )


        page_words = get_words(
            page_text
        )


        # ----------------------------------------------------
        # Keyword overlap
        # ----------------------------------------------------

        keyword_matches = (

            content_words.intersection(
                page_words
            )

        )


        if content_words:

            keyword_score = (

                len(keyword_matches)

                /

                len(content_words)

            )

        else:

            keyword_score = 0.0


        # ----------------------------------------------------
        # Factual term overlap
        # ----------------------------------------------------

        factual_matches = (

            requested_factual_terms.intersection(
                page_words
            )

        )


        if requested_factual_terms:

            factual_score = (

                len(factual_matches)

                /

                len(requested_factual_terms)

            )

        else:

            factual_score = 0.0


        # ----------------------------------------------------
        # Exact phrase
        # ----------------------------------------------------

        phrase_score = 0.0


        if (

            len(question_normalized) >= 5

            and

            question_normalized in page_normalized

        ):

            phrase_score = 1.0


        # ----------------------------------------------------
        # Numeric / technical value overlap
        # ----------------------------------------------------

        page_values = set(

            re.findall(

                r"\b\d+(?:\.\d+){0,3}\b",

                page_text

            )

        )


        value_matches = (

            question_values.intersection(
                page_values
            )

        )


        if question_values:

            value_score = (

                len(value_matches)

                /

                len(question_values)

            )

        else:

            value_score = 0.0


        # ----------------------------------------------------
        # Important exact-term bonus
        # ----------------------------------------------------

        exact_term_bonus = 0.0


        for term in content_words:

            if term in page_words:

                exact_term_bonus += 1.0


        if content_words:

            exact_term_bonus = min(

                exact_term_bonus
                /

                len(content_words),

                1.0

            )


        # ----------------------------------------------------
        # Final retrieval score
        # ----------------------------------------------------

        score = (

            0.50 * keyword_score

            +

            0.25 * factual_score

            +

            0.10 * phrase_score

            +

            0.10 * value_score

            +

            0.05 * exact_term_bonus

        )


        page_scores.append(
            score
        )


    # ========================================================
    # CONVERT TO NUMPY ARRAY
    # ========================================================

    hybrid_scores = np.array(

        page_scores,

        dtype=float

    )


    # ========================================================
    # SELECT TOP PAGES
    # ========================================================

    top_k = min(

        3,

        len(text_pages)

    )


    top_indices = np.argsort(

        hybrid_scores

    )[::-1][:top_k]


    # ========================================================
    # BUILD FOCUSED EVIDENCE
    # ========================================================

    evidence_parts = []


    for index in top_indices:

        page = text_pages[index]

        page_text = page.get(
            "text",
            ""
        )


        # ----------------------------------------------------
        # Break page into useful units
        # ----------------------------------------------------

        units = split_into_units(

            page_text

        )


        scored_units = []


        for unit in units:

            unit_words = get_words(
                unit
            )


            # ------------------------------------------------
            # Keyword relevance
            # ------------------------------------------------

            unit_keyword_matches = (

                content_words.intersection(
                    unit_words
                )

            )


            if content_words:

                unit_keyword_score = (

                    len(
                        unit_keyword_matches
                    )

                    /

                    len(content_words)

                )

            else:

                unit_keyword_score = 0.0


            # ------------------------------------------------
            # Factual relevance
            # ------------------------------------------------

            unit_factual_matches = (

                requested_factual_terms.intersection(
                    unit_words
                )

            )


            if requested_factual_terms:

                unit_factual_score = (

                    len(
                        unit_factual_matches
                    )

                    /

                    len(
                        requested_factual_terms
                    )

                )

            else:

                unit_factual_score = 0.0


            # ------------------------------------------------
            # Numeric relevance
            # ------------------------------------------------

            unit_values = set(

                re.findall(

                    r"\b\d+(?:\.\d+){0,3}\b",

                    unit

                )

            )


            unit_value_matches = (

                question_values.intersection(
                    unit_values
                )

            )


            if question_values:

                unit_value_score = (

                    len(
                        unit_value_matches
                    )

                    /

                    len(
                        question_values
                    )

                )

            else:

                unit_value_score = 0.0


            # ------------------------------------------------
            # Unit score
            # ------------------------------------------------

            unit_score = (

                0.55 * unit_keyword_score

                +

                0.30 * unit_factual_score

                +

                0.15 * unit_value_score

            )


            scored_units.append(

                (
                    unit_score,
                    unit

                )

            )


        # ----------------------------------------------------
        # Sort best evidence first
        # ----------------------------------------------------

        scored_units.sort(

            key=lambda item: item[0],

            reverse=True

        )


        # ----------------------------------------------------
        # Keep relevant units
        # ----------------------------------------------------

        selected_units = [

            unit

            for score, unit

            in scored_units

            if score > 0

        ][:10]


        # ----------------------------------------------------
        # Fallback
        # ----------------------------------------------------

        if not selected_units:

            selected_units = [

                unit

                for score, unit

                in scored_units[:5]

            ]


        # ----------------------------------------------------
        # Page evidence
        # ----------------------------------------------------

        focused_text = "\n".join(

            selected_units

        )


        evidence_parts.append(

            f"PAGE {page['page']}:\n"
            f"{focused_text}"

        )


    # ========================================================
    # COMBINE EVIDENCE
    # ========================================================

    evidence = "\n\n".join(

        evidence_parts

    )


    # ========================================================
    # RETURN SAME STRUCTURE AS BEFORE
    # ========================================================

    return (

        hybrid_scores,

        top_indices,

        evidence,

        content_words,

        requested_factual_terms

    )


# ============================================================
# ============================================================
# FILE UPLOADER
# ============================================================

uploaded_file = st.file_uploader(

    "Upload a document or image",

    type=[

        "pdf",
        "docx",

        "png",
        "jpg",
        "jpeg"

    ]

)


# ============================================================
# PROCESS UPLOADED FILE
# ============================================================

if uploaded_file:

    os.makedirs(
        TEMP_DIR,
        exist_ok=True
    )


    extension = os.path.splitext(

        uploaded_file.name

    )[1].lower()


    file_path = os.path.join(

        TEMP_DIR,

        "uploaded" + extension

    )


    # --------------------------------------------------------
    # Save upload
    # --------------------------------------------------------

    with open(

        file_path,

        "wb"

    ) as f:

        f.write(

            uploaded_file.getbuffer()

        )


    # ========================================================
    # PDF
    # ========================================================

    if extension == ".pdf":

        pages = process_pdf(
            file_path
        )

        document_type = "PDF"


    # ========================================================
    # DOCX
    # ========================================================

    elif extension == ".docx":

        pages = process_docx(
            file_path
        )

        document_type = "DOCX"


    # ========================================================
    # IMAGE
    # ========================================================

    elif extension in [

        ".png",
        ".jpg",
        ".jpeg"

    ]:

        pages = process_image(
            file_path
        )

        document_type = "IMAGE"


    else:

        st.error(
            "Unsupported file type."
        )

        st.stop()


    # ========================================================
    # DOCUMENT STATUS
    # ========================================================

    text_pages = [

        page

        for page in pages

        if page["text"].strip()

    ]


    st.success(

        f"{document_type} loaded successfully — "
        f"{len(pages)} page(s)"

    )


    st.info(

        f"Readable text available on "
        f"{len(text_pages)} page(s)."

    )


    # ========================================================
    # IMAGE OCR DISPLAY
    # ========================================================

    if document_type == "IMAGE":

        st.subheader(
            "🖼️ Visual Evidence"
        )


        st.image(

            pages[0]["image"],

            caption="Uploaded image",

            use_container_width=True

        )


        st.subheader(
            "🔤 OCR Extracted Text"
        )


        if pages[0]["text"]:

            st.text_area(

                "Recognized text",

                pages[0]["text"],

                height=250

            )

        else:

            st.warning(

                "No readable text was detected."

            )


    # ========================================================
    # QUESTION INPUT
    # ========================================================

    question = st.text_input(

        "🔎 Ask a question about the document"

    )


    if question:

        # ====================================================
        # NO TEXT
        # ====================================================

        if not text_pages:

            st.warning(

                "No readable text was found "
                "for retrieval."

            )

            st.stop()


        # ====================================================
        # RETRIEVE EVIDENCE
        # ====================================================

        (
            hybrid_scores,
            top_indices,
            evidence,
            content_words,
            requested_factual_terms

        ) = retrieve_evidence(

            question,

            text_pages

        )


        # ====================================================
        # STRICT GROUNDING CHECK
        # ====================================================

        evidence_words = get_words(
            evidence
        )


        matched_content_words = (

            content_words.intersection(
                evidence_words
            )

        )


        grounded = (

            len(
                matched_content_words
            )

            >=

            1

        )


        # ----------------------------------------------------
        # Factual questions need factual evidence.
        # ----------------------------------------------------

        if requested_factual_terms:

            matched_factual_terms = (

                requested_factual_terms.intersection(
                    evidence_words
                )

            )


            factual_grounded = (

                len(
                    matched_factual_terms
                )

                >=

                1

            )


            grounded = (

                grounded

                and

                factual_grounded

            )


        # ====================================================
        # NOT FOUND
        # ====================================================

        if not grounded:

            st.divider()


            st.subheader(
                "🤖 Answer"
            )


            st.success(

                "Not found in the provided evidence."

            )


            st.write(

                "📌 **Primary supporting source:** "
                "No supporting source found"

            )


            st.subheader(
                "📚 Source Evidence"
            )


            for rank, index in enumerate(

                top_indices

            ):

                page = text_pages[index]


                with st.expander(

                    f"Result {rank + 1} — "
                    f"Page {page['page']}"

                ):

                    st.write(

                        f"**Retrieval relevance:** "
                        f"{hybrid_scores[index]:.3f}"

                    )


                    st.write(
                        page["text"]
                    )


                    if os.path.exists(
                        page["image"]
                    ):

                        st.image(

                            page["image"],

                            caption=(

                                f"Visual evidence — "
                                f"Page {page['page']}"

                            ),

                            use_container_width=True

                        )


            st.stop()


        # ====================================================
        # GENERATE GROUNDED ANSWER
        # ====================================================

        st.divider()


        st.subheader(
            "🤖 Answer"
        )


        try:

            with st.spinner(

                "Generating grounded answer..."

            ):

                answer = generate_answer(

                    question,

                    evidence

                )


            # ------------------------------------------------
            # Clean model output
            # ------------------------------------------------

            answer = re.sub(

                r"ANSWER:\s*",

                "",

                answer,

                flags=re.IGNORECASE

            )


            answer = re.sub(

                r"SOURCE_PAGE:\s*.*",

                "",

                answer,

                flags=re.IGNORECASE

            )


            answer = answer.strip()


            if not answer:

                answer = (

                    "Not found in the provided evidence."

                )


            # ------------------------------------------------
            # If Qwen itself returns the not-found response
            # ------------------------------------------------

            st.success(
                answer
            )


        except Exception as e:

            st.error(

                "Could not generate answer: "

                + str(e)

            )

            answer = ""


        # ====================================================
        # PRIMARY SUPPORTING SOURCE
        # ====================================================

        if answer:

            answer_words = get_words(
                answer
            )


            source_scores = []


            for index in top_indices:

                page_words = get_words(

                    text_pages[index]["text"]

                )


                answer_overlap = len(

                    answer_words.intersection(
                        page_words
                    )

                )


                query_overlap = len(

                    content_words.intersection(
                        page_words
                    )

                )


                source_score = (

                    0.70 * answer_overlap

                    +

                    0.30 * query_overlap

                )


                source_scores.append(
                    source_score
                )


            best_position = int(

                np.argmax(
                    source_scores
                )

            )


            best_index = (

                top_indices[
                    best_position
                ]

            )


            best_page = (

                text_pages[
                    best_index
                ]

            )


            st.write(

                f"📌 **Primary supporting "
                f"source: Page "
                f"{best_page['page']}**"

            )


        # ====================================================
        # SOURCE EVIDENCE
        # ====================================================

        st.subheader(
            "📚 Source Evidence"
        )


        for rank, index in enumerate(

            top_indices

        ):

            page = text_pages[index]


            with st.expander(

                f"Result {rank + 1} — "
                f"Page {page['page']}"

            ):

                st.write(

                    f"**Retrieval relevance:** "
                    f"{hybrid_scores[index]:.3f}"

                )


                col1, col2 = st.columns(
                    2
                )


                # ------------------------------------------------
                # Retrieved text
                # ------------------------------------------------

                with col1:

                    st.write(
                        "#### 📄 Retrieved Content"
                    )


                    if page["text"]:

                        st.write(
                            page["text"]
                        )

                    else:

                        st.write(
                            "No text available."
                        )


                # ------------------------------------------------
                # Visual evidence
                # ------------------------------------------------

                with col2:

                    st.write(
                        "#### 🖼️ Visual Evidence"
                    )


                    if os.path.exists(
                        page["image"]
                    ):

                        st.image(

                            page["image"],

                            caption=(

                                f"Page "
                                f"{page['page']}"

                            ),

                            use_container_width=True

                        )

                    else:

                        st.write(
                            "No visual preview available."
                        )